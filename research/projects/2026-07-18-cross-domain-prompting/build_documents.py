#!/usr/bin/env python3
"""Build the two Word deliverables for the cross-domain prompting study.

The source of truth remains the adjacent Markdown files.  This renderer applies
the repository's chosen document presets explicitly so the output is stable and
auditable rather than dependent on a workstation's Normal template.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUTPUT_DIR = HERE / "deliverables"
SKILL_DIR = Path(
    "/Users/kevinconnolly/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.715.12143/skills/documents"
)
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEXT = "20252A"
MUTED = "59636E"
PALE_BLUE = "EAF2F8"
PALE_GRAY = "F2F4F7"
MID_GRAY = "D7DDE3"
WHITE = "FFFFFF"


PRESETS = {
    "standard_business_brief": {
        "body_after": 6,
        "body_line": 1.10,
        "h1": (16, BLUE, 16, 8),
        "h2": (13, BLUE, 12, 6),
        "h3": (12, DARK_BLUE, 8, 4),
        "list_marker": 0.25,
        "list_text": 0.50,
        "list_hanging": 0.25,
        "list_after": 8,
        "list_line": 1.167,
        "table_header": PALE_GRAY,
        "table_font": 9.2,
        "code_font": 8.0,
    },
    "compact_reference_guide": {
        "body_after": 6,
        "body_line": 1.25,
        "h1": (16, BLUE, 18, 10),
        "h2": (13, BLUE, 14, 7),
        "h3": (12, DARK_BLUE, 10, 5),
        "list_marker": 0.187,
        "list_text": 0.375,
        "list_hanging": 0.188,
        "list_after": 4,
        "list_line": 1.25,
        "table_header": "E8EEF5",
        "table_font": 9.0,
        "code_font": 7.8,
    },
}


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    tr_pr.append(marker)


def set_shading(element, fill: str) -> None:
    shd = element.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        element.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_border(paragraph, side: str, color: str, size: int = 10, space: int = 4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    existing = borders.find(qn(f"w:{side}"))
    if existing is not None:
        borders.remove(existing)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def set_cell_border(cell, **sides) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side, attrs in sides.items():
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_style_font(style, name: str, size: float, color: str = TEXT, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)


def set_run_font(run, name: str = "Calibri", size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)


def configure_styles(doc: Document, preset_name: str) -> None:
    preset = PRESETS[preset_name]
    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11)
    normal.paragraph_format.space_after = Pt(preset["body_after"])
    normal.paragraph_format.line_spacing = preset["body_line"]
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    set_style_font(title, "Calibri", 23, TEXT, True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(7)
    title.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    set_style_font(caption, "Calibri", 7.5, MUTED, True)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(2)

    for style_name, key in (("Heading 1", "h1"), ("Heading 2", "h2"), ("Heading 3", "h3")):
        size, color, before, after = preset[key]
        style = doc.styles[style_name]
        set_style_font(style, "Calibri", size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        set_style_font(style, "Calibri", 11)
        style.paragraph_format.space_after = Pt(preset["list_after"])
        style.paragraph_format.line_spacing = preset["list_line"]
        style.paragraph_format.widow_control = True


def configure_page(doc: Document, running_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("PROMPTS FOR PROGRESS")
    left.bold = True
    left.font.color.rgb = RGBColor.from_string(BLUE)
    set_run_font(left, size=8)
    p.add_run("\t")
    right = p.add_run(running_title.upper())
    right.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_font(right, size=8)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    label = fp.add_run("PROMPTS FOR PROGRESS  ·  ")
    label.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_font(label, size=8)
    page_run = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_end)
    set_run_font(page_run, size=8)


def add_masthead(doc: Document, kicker: str, title: str, subtitle: str, metadata: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(kicker.upper())
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    set_run_font(run, size=8.5)

    p = doc.add_paragraph(style="Title")
    run = p.add_run(normalize_text(title))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(13)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(subtitle)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_font(run, size=12.5)

    for row_index in range(0, len(metadata), 2):
        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(3)
        p_meta.paragraph_format.tab_stops.add_tab_stop(Inches(3.25))
        pairs = metadata[row_index : row_index + 2]
        for pair_index, (label, value) in enumerate(pairs):
            if pair_index:
                p_meta.add_run("\t")
            label_run = p_meta.add_run(label.upper() + "  ")
            label_run.bold = True
            label_run.font.color.rgb = RGBColor.from_string(BLUE)
            set_run_font(label_run, size=8.5)
            value_run = p_meta.add_run(value)
            value_run.font.color.rgb = RGBColor.from_string(TEXT)
            set_run_font(value_run, size=9.5)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(5)
    rule.paragraph_format.space_after = Pt(12)
    set_paragraph_border(rule, "bottom", BLUE, size=14, space=1)


def normalize_text(text: str) -> str:
    replacements = {
        r"\mathbb F_3": "F₃",
        r"\subseteq": "⊆",
        r"\notin": "∉",
        r"\ge": "≥",
        r"\le": "≤",
        r"\times": "×",
        r"\to": "→",
        r"\neq": "≠",
        r"\in": "∈",
        r"\pmod": " mod",
        r"\sum": "Σ",
    }
    text = text.replace(r"\(", "").replace(r"\)", "")
    text = text.replace(r"\[", "").replace(r"\]", "")
    text = text.replace(r"\{", "{").replace(r"\}", "}")
    for old, new in replacements.items():
        text = text.replace(old, new)
    superscript = str.maketrans("0123456789+-abkn", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻ᵃᵇᵏⁿ")
    text = re.sub(
        r"\^\{([^{}]+)\}",
        lambda match: match.group(1).translate(superscript),
        text,
    )
    text = re.sub(
        r"\^([0-9abkn+-]+)",
        lambda match: match.group(1).translate(superscript),
        text,
    )
    text = text.replace("$", "")
    return text


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|(?<!\*)\*[^*]+\*(?!\*))")


def add_inline(paragraph, text: str, *, size: float | None = None) -> None:
    text = normalize_text(text)
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_run_font(run, size=size)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Courier New", size=(size or 10.5) - 0.5)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            set_run_font(run, size=size)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size)


def next_ids(doc: Document) -> tuple[int, int]:
    root = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in root.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in root.findall(qn("w:num"))]
    return (max(abstract_ids, default=-1) + 1, max(num_ids, default=0) + 1)


def add_numbering_definition(doc: Document, ordered: bool, preset_name: str) -> int:
    preset = PRESETS[preset_name]
    abstract_id, num_id = next_ids(doc)
    root = doc.part.numbering_part.element

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multilevel = OxmlElement("w:multiLevelType")
    multilevel.set(qn("w:val"), "singleLevel")
    abstract.append(multilevel)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(round(preset["list_text"] * 1440)))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(round(preset["list_text"] * 1440)))
    ind.set(qn("w:hanging"), str(round(preset["list_hanging"] * 1440)))
    p_pr.append(ind)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), "Calibri")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    root.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    root.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_ref)
    p_pr.append(num_pr)


def add_code_block(doc: Document, lines: list[str], preset_name: str, language: str) -> None:
    preset = PRESETS[preset_name]
    first_line = lines[0].strip() if lines else ""
    if first_line.startswith('"""Self-contained 512-cap'):
        # Keep the prose instruction that introduces these long, self-contained
        # blocks with the code instead of leaving it stranded on the prior page.
        if doc.paragraphs:
            intro = doc.paragraphs[-1]
            intro.paragraph_format.page_break_before = True
            intro.paragraph_format.keep_with_next = True
        else:
            page_break = doc.add_paragraph()
            page_break.add_run().add_break(WD_BREAK.PAGE)
    if language:
        label = doc.add_paragraph(style="Caption")
        label.paragraph_format.keep_with_next = True
        r = label.add_run(language.upper())
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(MUTED)
        set_run_font(r, size=7.5)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.widow_control = False
    p.paragraph_format.keep_together = len(lines) <= 35
    set_shading(p._p.get_or_add_pPr(), PALE_GRAY)
    set_paragraph_border(p, "left", BLUE, size=10, space=5)
    for index, line in enumerate(lines or [""]):
        run = p.add_run(line)
        set_run_font(run, "Courier New", preset["code_font"])
        if index != len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def parse_table_row(line: str) -> list[str]:
    text = line.strip().strip("|")
    return [cell.strip() for cell in text.split("|")]


def is_separator_row(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown_table(doc: Document, rows: list[list[str]], preset_name: str) -> None:
    if not rows:
        return
    preset = PRESETS[preset_name]
    col_count = max(len(row) for row in rows)
    normalized = [row + [""] * (col_count - len(row)) for row in rows]
    maxima = [max(8, min(34, max(len(row[c]) for row in normalized))) for c in range(col_count)]
    widths = column_widths_from_weights(maxima, 9360)

    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.keep_with_next = True

    table = doc.add_table(rows=len(normalized), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)
    for r_index, values in enumerate(normalized):
        row = table.rows[r_index]
        prevent_row_split(row)
        if r_index == 0:
            set_repeat_table_header(row)
        for c_index, value in enumerate(values):
            cell = row.cells[c_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            cell_p = cell.paragraphs[0]
            cell_p.paragraph_format.space_after = Pt(2.5)
            cell_p.paragraph_format.line_spacing = 1.0
            add_inline(cell_p, value, size=preset["table_font"])
            if r_index == 0:
                set_shading(cell._tc.get_or_add_tcPr(), preset["table_header"])
                for run in cell_p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            border = {"val": "single", "sz": "4", "color": MID_GRAY}
            set_cell_border(cell, top=border, left=border, bottom=border, right=border)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_together = True
    set_shading(p._p.get_or_add_pPr(), PALE_BLUE)
    set_paragraph_border(p, "left", BLUE, size=14, space=6)
    add_inline(p, text)
    for run in p.runs:
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)


def markdown_to_doc(doc: Document, path: Path, preset_name: str) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    # The Markdown H1 and dateline are represented by the Word masthead.
    if lines and lines[0].startswith("# "):
        index = 1
    while index < len(lines) and (not lines[index].strip() or lines[index].startswith("**")):
        index += 1

    list_kind: str | None = None
    list_num_id: int | None = None

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            list_kind = None
            list_num_id = None
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            add_code_block(doc, code_lines, preset_name, language)
            list_kind = None
            list_num_id = None
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_separator_row(lines[index + 1]):
            rows = [parse_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(parse_table_row(lines[index]))
                index += 1
            add_markdown_table(doc, rows, preset_name)
            list_kind = None
            list_num_id = None
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1)) - 1
            text_value = normalize_text(heading_match.group(2))
            if "FULL PROMPT — BEGIN COPY" in text_value:
                p_break = doc.add_paragraph()
                p_break.add_run().add_break(WD_BREAK.PAGE)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, text_value)
            if "FULL PROMPT" in text_value:
                set_shading(p._p.get_or_add_pPr(), PALE_BLUE)
                set_paragraph_border(p, "left", BLUE, size=14, space=4)
            list_kind = None
            list_num_id = None
            index += 1
            continue

        if stripped == "---":
            rule = doc.add_paragraph()
            rule.paragraph_format.space_before = Pt(4)
            rule.paragraph_format.space_after = Pt(7)
            set_paragraph_border(rule, "bottom", MID_GRAY, size=6, space=1)
            list_kind = None
            list_num_id = None
            index += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            add_callout(doc, " ".join(quote_lines))
            list_kind = None
            list_num_id = None
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        bulleted = re.match(r"^[-*]\s+(.+)$", stripped)
        if numbered or bulleted:
            kind = "ordered" if numbered else "bullet"
            if kind != list_kind or list_num_id is None:
                list_num_id = add_numbering_definition(doc, ordered=(kind == "ordered"), preset_name=preset_name)
                list_kind = kind
            p = doc.add_paragraph(style="List Number" if kind == "ordered" else "List Bullet")
            apply_num(p, list_num_id)
            add_inline(p, (numbered or bulleted).group(1))
            index += 1
            continue

        # Consecutive non-structural Markdown lines are one paragraph.
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (
                candidate.startswith(("#", "```", ">", "|"))
                or candidate == "---"
                or re.match(r"^\d+\.\s+", candidate)
                or re.match(r"^[-*]\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))
        list_kind = None
        list_num_id = None


def add_end_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    set_paragraph_border(p, "top", MID_GRAY, size=6, space=6)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_font(r, size=8.5)


def build_document(
    *,
    source: str,
    output: str,
    preset: str,
    kicker: str,
    title: str,
    subtitle: str,
    running_title: str,
    metadata: list[tuple[str, str]],
    subject: str,
    include_end_note: bool = True,
) -> Path:
    doc = Document()
    configure_styles(doc, preset)
    configure_page(doc, running_title)
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Prompts for Progress"
    props.keywords = "prompting research, research agents, verification, prompt design"
    props.comments = "Generated from repository Markdown; layout verified before delivery."

    add_masthead(doc, kicker, title, subtitle, metadata)
    markdown_to_doc(doc, HERE / source, preset)
    if include_end_note:
        add_end_note(doc, "Repository snapshot: 18 July 2026  ·  Prompts for Progress")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    target = OUTPUT_DIR / output
    doc.save(target)
    return target


def main() -> None:
    report = build_document(
        source="cross-domain-prompting-style-research.md",
        output="cross-domain-prompting-style-research.docx",
        preset="standard_business_brief",
        kicker="Research note 01 · Cross-domain corpus study",
        title="Cross-Domain Prompting Styles for Research Progress",
        subtitle="A descriptive analysis of how research prompts turn uncertainty into checkable state",
        running_title="Cross-domain prompting styles",
        metadata=[
            ("Corpus", "23 manifests; 21 content-bearing"),
            ("Snapshot", "18 July 2026"),
            ("Evidence", "Descriptive; non-causal"),
            ("Companion", "Cap-set 513 prompt dossier"),
        ],
        subject="Cross-domain prompt-method analysis of the Prompts for Progress archive",
    )
    prompt = build_document(
        source="cap-set-513-research-prompt.md",
        output="cap-set-513-research-prompt.docx",
        preset="compact_reference_guide",
        kicker="Research prompt 01 · Executable combinatorics",
        title="Exceed the Archived 512-Cap in F₃⁸",
        subtitle="A self-contained, verifier-gated prompt for a fresh research session",
        running_title="Cap-set 513 prompt dossier",
        metadata=[
            ("Target", "Valid cap of size ≥ 513 in F₃⁸"),
            ("Baseline", "Valid archived cap of size 512"),
            ("Verification", "Python standard library"),
            ("Novelty", "Archive-relative until checked"),
        ],
        subject="Standalone research prompt for an independently verifiable cap-set construction",
        include_end_note=False,
    )
    print(report)
    print(prompt)


if __name__ == "__main__":
    main()
